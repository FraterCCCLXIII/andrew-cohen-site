#!/usr/bin/env python3
"""Extract additional Andrew Cohen books for the online reader."""

from __future__ import annotations

import re
import sys
from pathlib import Path

import fitz
from docx import Document
from docx.oxml.ns import qn

sys.path.insert(0, str(Path(__file__).resolve().parent))
from lib.book_extract import (  # noqa: E402
    BOOK_SOURCES,
    ROOT,
    extract_body,
    pages_text,
    slugify,
    unique_id,
    write_book,
)

FREEDOM_CHAPTERS = [
    ("foreword", "Foreword", "front-matter", 4, 6),
    ("introduction", "Introduction", "front-matter", 7, 8),
    ("preface", "Preface", "front-matter", 9, 9),
    (1, "An Overwhelming Passion", "teachings", 10, 11),
    (2, "One without a Second", "teachings", 12, 14),
    (3, "A Stand That Is Immovable", "teachings", 15, 16),
    (4, "Abandon the Future", "teachings", 17, 19),
    (5, "Cynicism", "teachings", 20, 21),
    (6, "What Am I Doing?", "teachings", 22, 23),
    (7, "An Absolute Choice", "teachings", 24, 24),
    (8, "A Perfect Relationship with Thought", "teachings", 25, 27),
    (9, "Stop Struggling", "teachings", 28, 30),
    (10, "Hold onto Nothing", "teachings", 31, 32),
    (11, "The Impersonal View", "teachings", 33, 34),
    (12, "True Finders Are Few", "teachings", 35, 35),
    (13, "The Edge", "teachings", 36, 38),
    (14, "Like a Twig in a Hurricane", "teachings", 39, 40),
    (15, "Perfect Aloneness", "teachings", 41, 42),
    (16, "A Trust That Is Absolute", "teachings", 43, 43),
    (17, "Integrity Is Multidimensional", "teachings", 44, 45),
    (18, "Spiritual Conscience", "teachings", 46, 47),
    (19, "This Mysterious Secret", "teachings", 48, 50),
    (20, "The Personalization of Feeling", "teachings", 51, 52),
    (21, "A Love So Great", "teachings", 53, 55),
    (22, "Ecstatic Intimacy", "teachings", 56, 57),
    (23, "An Immaculate Condition", "teachings", 58, 59),
    (24, "Final Purity", "teachings", 60, 60),
    (25, "What Would You Do?", "teachings", 61, 63),
]

LIVING_CHAPTERS = [
    ("foreword", "Foreword by Ken Wilber", "front-matter", 13, 20),
    ("preface", "Preface", "front-matter", 21, 24),
    (1, "What Is Enlightenment?", "dialogues", 25, 30),
    (2, "That Which Was Never Born", "dialogues", 31, 36),
    (3, "The Perfect Response", "dialogues", 37, 42),
    (4, "One Without a Second", "dialogues", 43, 46),
    (5, "The Status Quo Cannot Contain It", "dialogues", 47, 54),
    (6, "Revelation and Awakening", "dialogues", 55, 64),
    (7, "To Cast No Shadow", "dialogues", 65, 70),
    (8, "The Ultimate Challenge", "dialogues", 71, 76),
    (9, "A Straight Line to the Absolute", "dialogues", 77, 82),
    (10, "Grace Is Not Enough", "dialogues", 83, 86),
    (11, "Mere Mortals", "dialogues", 87, 92),
    (12, "The End of Karma", "dialogues", 93, 98),
    (13, "True Conscience", "dialogues", 99, 104),
    (14, "The Only Obstacle", "dialogues", 105, 112),
    (15, "Liberation Beyond Gender", "dialogues", 113, 118),
    (16, "The Promise of Perfection", "dialogues", 119, 126),
    (17, "An Unbroken Universal Unfolding", "dialogues", 127, 132),
    (18, "Impersonal Enlightenment", "dialogues", 133, 138),
    (19, "The Complete Picture", "dialogues", 139, 144),
    (20, "The Imperative to Evolve", "dialogues", 145, 153),
]

THREE_JEWELS_SECTIONS = [
    ("seductive-allure", "The Seductive Allure of Spiritual Awakening", "triple-gem", 6, 7),
    ("sacred-synergy", "The Sacred Synergy", "triple-gem", 8, 11),
    ("brahman-is-the-world", "Brahman is the World", "dharma", 12, 13),
    (
        "evolutionary-enlightenment-is-the-new-enlightenment",
        "Evolutionary Enlightenment is the New Enlightenment",
        "dharma",
        14,
        17,
    ),
    ("god-evolves-through-us", "God Evolves Through Us", "dharma", 18, 20),
    ("how-shall-i-live", "How Shall I Live? — Awakening Your Authentic Self", "dharma", 21, 23),
    (
        "being-evolutionarily-enlightened",
        "Being Evolutionarily Enlightened: What Does it Really Mean?",
        "dharma",
        24,
        27,
    ),
    ("whole-of-the-holy-life", "The Whole of the Holy Life", "sangha", 28, 29),
    ("living-awakening-relationally", "Living Awakening Relationally", "sangha", 30, 30),
    ("delight-intersubjective-dialogue", "The Delight of Intersubjective Dialogue", "sangha", 31, 34),
    (
        "power-enlightened-communication",
        "The Power of Enlightened Communication in our World in Crisis",
        "sangha",
        35,
        37,
    ),
    ("unique-role-of-the-guru", "The Unique Role of the Guru", "buddha", 38, 39),
    ("being-the-face-of-your-true-self", "Being the Face of your own True Self", "buddha", 40, 40),
    ("being-a-clear-mirror", "Being a clear mirror, reflecting all of who you are", "buddha", 41, 41),
    (
        "shining-a-light",
        "Shining a light on where we least want to look…",
        "buddha",
        42,
        42,
    ),
    (
        "transmitting-enlightened-awareness",
        "Transmitting enlightened awareness directly to you",
        "buddha",
        43,
        43,
    ),
    ("the-man-and-the-master", "The Man and the Master", "buddha", 44, 47),
    ("co-creating-manifest-nirvana", "Co-creating Manifest Nirvana", "manifest-nirvana", 48, 51),
]


def extract_pdf_book(
    *,
    slug: str,
    pdf_name: str,
    title: str,
    subtitle: str,
    credits: str,
    parts: list[dict[str, str]],
    sections: list[tuple],
) -> None:
    pdf = BOOK_SOURCES / pdf_name
    doc = fitz.open(pdf)
    chapters: list[dict] = []
    existing: set[str] = set()

    for entry in sections:
        key, chapter_title, part_id, start, end = entry
        chapter_id = key if isinstance(key, str) else f"chapter-{key}"

        raw = pages_text(doc, start, end)
        body = extract_body(raw, chapter_title)
        if len(body) < 60 and chapter_id not in {"foreword", "preface", "introduction", "acknowledgements"}:
            print(f"Warning: short chapter {chapter_title} ({len(body)} chars)")

        if chapter_id in existing:
            chapter_id = unique_id(chapter_title, existing)
        else:
            existing.add(chapter_id)

        chapters.append(
            {
                "id": chapter_id,
                "title": chapter_title,
                "partId": part_id,
                "body": body,
            }
        )

    write_book(
        slug=slug,
        title=title,
        subtitle=subtitle,
        credits=credits,
        parts=parts,
        chapters=chapters,
    )


def extract_freedom() -> None:
    extract_pdf_book(
        slug="freedom-has-no-history",
        pdf_name="freedom-has-no-history.pdf",
        title="Freedom Has No History",
        subtitle="A Call to Awaken",
        credits="Copyright © 2024 by Andrew Cohen · First Edition 1997 · ISBN 1-883929-17-2 · Published by Manifest Nirvana",
        parts=[
            {"id": "front-matter", "label": "Front Matter"},
            {"id": "teachings", "label": "Teachings"},
        ],
        sections=FREEDOM_CHAPTERS,
    )


def extract_living() -> None:
    extract_pdf_book(
        slug="living-enlightenment",
        pdf_name="living-enlightenment.pdf",
        title="Living Enlightenment",
        subtitle="A Call for Evolution Beyond Ego",
        credits="Copyright © 2002 by Moksha Press · ISBN 0-9653601-8-2",
        parts=[
            {"id": "front-matter", "label": "Front Matter"},
            {"id": "dialogues", "label": "Dialogues"},
        ],
        sections=LIVING_CHAPTERS,
    )


def extract_three_jewels() -> None:
    extract_pdf_book(
        slug="the-three-jewels-of-the-new-enlightenment",
        pdf_name="the-three-jewels-of-the-new-enlightenment.pdf",
        title="The Three Jewels of the New Enlightenment",
        subtitle="Reimagining the Buddha, Dharma, and Sangha for the 21st Century",
        credits="By Andrew Cohen and Hans Plasqui · Edited and compiled by Hans Plasqui",
        parts=[
            {"id": "triple-gem", "label": "The Power of the Triple Gem"},
            {"id": "dharma", "label": "The Dharma — Evolutionary Enlightenment"},
            {"id": "sangha", "label": "The Sangha — Intersubjective Nonduality"},
            {"id": "buddha", "label": "The Buddha — Portal to Your Own True Self"},
            {"id": "manifest-nirvana", "label": "Co-creating Manifest Nirvana"},
        ],
        sections=THREE_JEWELS_SECTIONS,
    )


def clean_docx_text(text: str) -> str:
    text = text.replace("\u00a0", " ")
    text = text.replace("\u2019", "'").replace("\u2018", "'")
    text = text.replace("\u201c", '"').replace("\u201d", '"')
    text = text.replace("\u2013", "-").replace("\u2014", "-")
    text = re.sub(r"\*+", "", text)
    text = re.sub(r"([A-Za-z])\.([A-Za-z])", r"\1. \2", text)
    text = re.sub(r"([A-Za-z])(\d)", r"\1. \2", text)
    text = re.sub(r"\.{3}", "...", text)
    text = re.sub(r"([A-Za-z])\s*\.\s*\.\s*\.\s*([A-Za-z])", r"\1.\2.", text)
    text = re.sub(r"([A-Za-z])\s*\.\s*([A-Za-z])\s*\.\s*([A-Za-z])\s*\.", r"\1.\2.\3.", text)
    text = re.sub(
        r"\b(January|February|March|April|May|June|July|August|September|October|November|December)(\d{1,2})",
        r"\1 \2",
        text,
    )
    text = re.sub(r"\bPart(One|Two|Three|Four)\b", r"Part \1", text, flags=re.I)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def paragraph_text(paragraph) -> str:
    return clean_docx_text(paragraph.text)


def paragraphs_to_body(paragraphs: list[str]) -> str:
    cleaned = [clean_docx_text(p) for p in paragraphs if clean_docx_text(p)]
    return "\n\n".join(cleaned)


def extract_my_master() -> None:
    docx_path = ROOT / "public/downloads/books/my-master-is-my-self.docx"
    if not docx_path.exists():
        raise FileNotFoundError(f"Missing source file: {docx_path}")

    doc = Document(docx_path)
    paragraphs = [p.text for p in doc.paragraphs]

    part_indices: list[tuple[str, int, str]] = []
    for i, raw in enumerate(paragraphs):
        text = clean_docx_text(raw)
        if re.fullmatch(r"Part (One|Two|Three|Four)", text, re.I):
            part_indices.append((text, i, slugify(text)))

    parts = [
        {"id": "front-matter", "label": "Front Matter"},
        {"id": "part-one", "label": "Part One"},
        {"id": "part-two", "label": "Part Two"},
        {"id": "part-three", "label": "Part Three"},
        {"id": "part-four", "label": "Part Four"},
    ]

    part_titles = {
        "part-one": "Part One · Diary March 25 – April 14, 1986",
        "part-two": "Part Two · Diary and Letters April – May 1986",
        "part-three": "Part Three · Letters May – July 1986",
        "part-four": "Part Four · Letters between Andrew and H.W.L. Poonja",
    }

    foreword_start = next(
        (i for i, raw in enumerate(paragraphs) if clean_docx_text(raw).startswith("Editor's Foreword")),
        29,
    )
    first_part_index = part_indices[0][1] if part_indices else len(paragraphs)

    chapters: list[dict] = []

    front_text = paragraphs_to_body(paragraphs[foreword_start:first_part_index])
    front_body = extract_body(front_text, "Front Matter")
    if len(front_body) > 40:
        chapters.append(
            {
                "id": "editors-foreword-and-introduction",
                "title": "Editor's Foreword and Introduction",
                "partId": "front-matter",
                "body": front_body,
            }
        )

    for idx, (_label, start, part_id) in enumerate(part_indices):
        end = part_indices[idx + 1][1] if idx + 1 < len(part_indices) else len(paragraphs)
        body = extract_body(paragraphs_to_body(paragraphs[start:end]), part_titles[part_id])
        chapters.append(
            {
                "id": part_id,
                "title": part_titles[part_id],
                "partId": part_id,
                "body": body,
            }
        )

    write_book(
        slug="my-master-is-my-self",
        title="My Master Is My Self",
        subtitle="The Birth of a Spiritual Teacher",
        credits="Copyright © 1989, 1995 by Moksha Foundation, Inc. · ISBN 1-883929-07-5",
        parts=parts,
        chapters=chapters,
    )


def convert_my_master_docx_to_pdf() -> None:
    docx_path = ROOT / "public/downloads/books/my-master-is-my-self.docx"
    pdf_path = ROOT / "public/downloads/books/my-master-is-my-self.pdf"
    doc = Document(docx_path)

    pdf = fitz.open()
    page = pdf.new_page(width=612, height=792)
    margin_x = 54
    y = 54
    line_height = 14
    page_width = 612 - margin_x * 2
    font = "helv"
    size = 11

    def new_page() -> None:
        nonlocal page, y
        page = pdf.new_page(width=612, height=792)
        y = 54

    def write_line(text: str, bold: bool = False) -> None:
        nonlocal y
        if y > 740:
            new_page()
        page.insert_text(
            (margin_x, y),
            text,
            fontname="hebo" if bold else font,
            fontsize=size if not bold else size + 1,
        )
        y += line_height + (4 if bold else 0)

    def write_wrapped(text: str, bold: bool = False) -> None:
        words = text.split()
        line = ""
        for word in words:
            candidate = f"{line} {word}".strip()
            if fitz.get_text_length(candidate, fontname="hebo" if bold else font, fontsize=size) > page_width:
                write_line(line, bold=bold)
                line = word
            else:
                line = candidate
        if line:
            write_line(line, bold=bold)

    skip_until_foreword = True
    for paragraph in doc.paragraphs:
        text = clean_docx_text(paragraph.text)
        if not text:
            y += 6
            continue
        if skip_until_foreword:
            if text.startswith("Editor's Foreword"):
                skip_until_foreword = False
            else:
                continue
        if re.fullmatch(r"Part (One|Two|Three|Four)", text, re.I):
            y += 10
            write_wrapped(text, bold=True)
            continue
        if re.fullmatch(r"Day (One|Two|Three|Four|Five|Six|Seven|Eight|Nine|Ten|\d+)", text, re.I):
            y += 8
            write_wrapped(text, bold=True)
            continue
        if paragraph.style and paragraph.style.name.startswith("Heading"):
            y += 8
            write_wrapped(text, bold=True)
            continue
        write_wrapped(text)

    pdf.save(pdf_path)
    print(f"Wrote PDF to {pdf_path}")


EXTRACTORS = {
    "freedom-has-no-history": extract_freedom,
    "living-enlightenment": extract_living,
    "the-three-jewels-of-the-new-enlightenment": extract_three_jewels,
    "my-master-is-my-self": extract_my_master,
}


def main() -> None:
    import argparse

    parser = argparse.ArgumentParser(description="Extract additional books")
    parser.add_argument("slug", nargs="?", help="Book slug to extract")
    parser.add_argument("--convert-my-master-pdf", action="store_true")
    args = parser.parse_args()

    if args.convert_my_master_pdf:
        convert_my_master_docx_to_pdf()
        return

    if args.slug:
        if args.slug not in EXTRACTORS:
            print(f"Unknown slug: {args.slug}")
            sys.exit(1)
        EXTRACTORS[args.slug]()
        return

    convert_my_master_docx_to_pdf()
    for slug, fn in EXTRACTORS.items():
        print(f"\n--- {slug} ---")
        fn()


if __name__ == "__main__":
    main()
